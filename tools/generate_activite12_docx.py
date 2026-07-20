from pathlib import Path
import csv
import sys

sys.path.insert(0, "/tmp/activite12-docx")
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / "livrables" / "admin-windows" / "activite12"
OUTDIR.mkdir(parents=True, exist_ok=True)
OUT = OUTDIR / "Dossier-Exploitation-SYS01b.docx"
CONVENTION = OUTDIR / "HIMBLOT-Olivier-Labo-Activite12-Dossier-Exploitation-SYS01b.docx"
IMG = ROOT / "docs/assets/img/admin-windows"

doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(1.8), Cm(1.7)
sec.left_margin, sec.right_margin = Cm(2), Cm(2)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
for n, size, color in [("Title", 26, "17365D"), ("Heading 1", 17, "17365D"), ("Heading 2", 13, "2F5597"), ("Heading 3", 11, "4472C4")]:
    styles[n].font.name = "Aptos Display"
    styles[n].font.size = Pt(size)
    styles[n].font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        hdr.cells[i].text = str(h)
        shade(hdr.cells[i], "D9EAF7")
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in hdr.cells[i].paragraphs[0].runs: r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(.5)
    run = p.add_run(text.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F2F2"); pPr.append(shd)

def proof(path, caption, context, result):
    full = IMG / path
    doc.add_paragraph(f"Contexte — {context}")
    if full.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(full), width=Cm(15.5))
        c = doc.add_paragraph(f"Preuve — {caption}")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs: r.italic = True
    else:
        doc.add_paragraph(f"Capture attendue non trouvée : {full.name}")
    doc.add_paragraph(f"Résultat interprété — {result}")

def page_break(): doc.add_page_break()

# En-tête / pied de page
header = sec.header.paragraphs[0]
header.text = "CORP.LOCAL — DOSSIER D’EXPLOITATION SYS01b"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor(100,100,100)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("HIMBLOT Olivier — Activité 12 — Interne / sans secrets     Page ")
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)

# Couverture
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DOSSIER D’EXPLOITATION\nSYS01b"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(23,54,93)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Environnement Windows — domaine corp.local"); r.font.size = Pt(16)
doc.add_paragraph("\n")
table(["Champ", "Valeur"], [
    ("Auteur", "HIMBLOT Olivier"), ("Site", "Labo"),
    ("Activité", "Activité 12 — Dossier exploitation"), ("Version", "1.0"),
    ("Date", "20 juillet 2026"), ("Classification", "Interne — aucun mot de passe ni clé en clair"),
])
doc.add_paragraph("Objet : permettre à un autre administrateur de comprendre, contrôler et reprendre l’environnement sans dépendre de connaissances implicites.")
doc.add_paragraph("Convention : HIMBLOT-Olivier-Labo-Activite12-Dossier-Exploitation-SYS01b.docx")
page_break()

doc.add_heading("Table des matières", level=1)
p = doc.add_paragraph()
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u'); p._p.append(fld)
doc.add_paragraph("Dans Word : clic droit sur la table > Mettre à jour les champs.")

doc.add_heading("1. Objet, périmètre et règles de sécurité", level=1)
doc.add_paragraph("Le dossier couvre LABO, SRV-AD01, SRV-FIC01 et POSTE-01, ainsi que l’annuaire, les GPO, la protection des postes, les fichiers, les restaurations, l’automatisation et le pare-feu.")
bullets([
    "Les valeurs issues d’exports et captures sont présentées comme constatées ; les valeurs seulement prévues par les procédures sont signalées comme cibles.",
    "Les secrets DSRM, BitLocker et LAPS ne figurent jamais dans ce document. Leur récupération doit suivre la procédure habilitée.",
    "Les captures sont accompagnées d’un contexte, d’un objectif de contrôle et d’une interprétation.",
    "Avant toute modification : sauvegarder, utiliser -WhatIf lorsque disponible, journaliser et prévoir le retour arrière.",
])
table(["Source de preuve", "Date/état", "Usage"], [
    ("AD-Inventory.html", "Export du 10/07/2026 13:37", "Comptes, groupes, ordinateurs, OU"),
    ("GPO-FW-*-Rules.csv", "Présent dans le module", "Inventaire de règles entrantes"),
    ("Captures it-1 à it-4", "Présentes dans le module", "Validation visuelle contextualisée"),
    ("Guides d’activités", "Configuration de référence", "Procédures et commandes de reprise"),
])

doc.add_heading("2. Présentation de l’environnement", level=1)
table(["Élément", "Adresse / OS", "Rôle", "État documentaire"], [
    ("LABO (labo_core)", "10.42.0.2/24 — Windows Server Core", "Hôte Hyper-V, Windows Admin Center", "IP constatée dans la fiche d’installation"),
    ("SRV-AD01", "10.42.0.10/24 — Windows Server 2025 Datacenter", "Contrôleur de domaine, AD DS, DNS", "IP fiche + objet AD exporté"),
    ("SRV-FIC01", "10.42.0.30/24 — Windows Server 2025 Datacenter", "Serveur membre, fichiers, VSS, sauvegarde", "IP confirmée + objet AD constaté"),
    ("POSTE-01", "10.42.0.20/24 — Windows 11 Pro", "Poste client joint au domaine", "IP confirmée + objet AD constaté"),
])
table(["Paramètre réseau", "Valeur"], [
    ("Réseau", "10.42.0.0/24"), ("Passerelle", "10.42.0.1"),
    ("DNS des membres", "10.42.0.10 (SRV-AD01)"), ("Domaine / NetBIOS", "corp.local / CORP"),
    ("Administration", "Windows Admin Center sur LABO, PowerShell Remoting et RDP si justifié"),
])
doc.add_paragraph("Dépendance critique : les membres du domaine doivent utiliser le DNS AD (SRV-AD01), pas un DNS public. Les DNS publics ne servent qu’en redirecteurs côté serveur DNS.")
proof("it-1/act3 Get-ADDomainController.png", "Contrôleur de domaine identifié", "Contrôler que SRV-AD01 porte le rôle attendu.", "La commande d’inventaire AD identifie SRV-AD01 comme contrôleur du domaine.")

doc.add_heading("3. Active Directory : OU, comptes et groupes", level=1)
doc.add_paragraph("Inventaire constaté : 9 utilisateurs, 58 groupes, 3 ordinateurs et 6 OU lors de l’export du 10/07/2026.")
table(["OU", "Contenu / finalité"], [
    ("Domain Controllers", "SRV-AD01 ; conteneur des contrôleurs de domaine"),
    ("Utilisateurs", "Comptes métiers du domaine"), ("Groupes", "Groupes globaux et locaux de domaine"),
    ("Ordinateurs", "POSTE-01 et futurs postes clients"), ("Serveurs", "SRV-FIC01 et futurs serveurs membres"),
    ("Administration", "Comptes et groupes d’administration dédiés"),
])
table(["Compte", "État", "OU", "Groupes métiers"], [
    ("user.rh1", "Actif", "Utilisateurs", "GG_RH"), ("user.rh2", "Actif", "Utilisateurs", "GG_RH"),
    ("user.it1", "Actif", "Utilisateurs", "GG_IT, GG_ADMIN"), ("user.rh3 (Alice Martin)", "Actif", "Utilisateurs", "GG_RH"),
    ("user.it2 (Bruno Durand)", "Actif", "Utilisateurs", "GG_IT"), ("user.rh4 (Claire Bernard)", "Actif", "Utilisateurs", "GG_RH"),
])
table(["Groupe", "Portée", "Fonction"], [
    ("GG_RH", "Global", "Comptes RH"), ("GG_IT", "Global", "Comptes IT"), ("GG_ADMIN", "Global", "Administration déléguée"),
    ("DL_RH_RW", "Local domaine", "Modification sur RH"), ("DL_IT_RW", "Local domaine", "Modification sur IT"),
    ("DL_COMMUN_RW", "Local domaine", "Modification sur COMMUN"),
])
doc.add_paragraph("Règle de gestion : créer un compte nominatif dans la bonne OU, l’affecter à un groupe global métier, ne jamais attribuer directement une ACL à un utilisateur, désactiver immédiatement les comptes sortants puis traiter leur suppression selon la politique de conservation.")
proof("it-2/OU principales.png", "Organisation des OU", "Vérifier visuellement l’arborescence logique du domaine.", "Les OU principales nécessaires à la délégation et au ciblage des GPO sont présentes.")

doc.add_heading("4. GPO et preuves d’application", level=1)
table(["GPO", "Cible", "Configuration / objectif", "Preuve à contrôler"], [
    ("GPO-PasswordPolicy", "Domaine", "Documentation de la politique ; efficacité à vérifier via Get-ADDefaultDomainPasswordPolicy", "Export GPO + commande AD"),
    ("GPO-Restriction-Panel", "OU Utilisateurs", "Interdire le panneau de configuration", "gpresult utilisateur"),
    ("GPO-Deploy-7Zip", "OU Ordinateurs", "Déploiement MSI depuis un chemin UNC", "7-Zip installé + gpresult"),
    ("GPO-BitLocker", "OU Ordinateurs", "XTS-AES 128, protecteur TPM/récupération, sauvegarde AD exigée", "manage-bde + objet de récupération masqué"),
    ("GPO-LAPS", "OU Ordinateurs", "Mot de passe local unique, rotation et sauvegarde AD", "journal LAPS + récupération masquée"),
    ("GPO-Map-Drives", "OU Utilisateurs", "Lecteurs H:, I:, S: selon groupes", "net use / lecteurs visibles"),
    ("GPO-FW-SRV-AD01", "OU Domain Controllers", "Pare-feu entrant adapté à AD/DNS", "CSV + tests de ports"),
    ("GPO-FW-SRV-FIC01", "OU Serveurs", "Pare-feu entrant adapté à SMB/admin", "CSV + tests de ports"),
])
code('''Get-GPO -All | Select DisplayName, GpoStatus
Get-GPInheritance -Target "OU=Ordinateurs,DC=corp,DC=local"
gpupdate /force
gpresult /h C:\\Temp\\gpresult.html
Get-ADDefaultDomainPasswordPolicy''')
doc.add_paragraph("Retour arrière GPO : désactiver le lien (sans supprimer immédiatement la GPO), forcer l’actualisation sur un poste pilote, vérifier gpresult et le journal GroupPolicy, puis documenter l’incident.")
proof("it-2/gpohtml.png", "Rapport de stratégie de groupe", "Prouver l’application effective des stratégies sur le poste, au-delà de leur simple existence dans GPMC.", "Un rapport gpresult HTML a été généré dans le module.")

doc.add_heading("5. BitLocker et Windows LAPS", level=1)
doc.add_heading("5.1 BitLocker", level=2)
bullets(["Cible : POSTE-01, Windows 11 Pro, VM génération 2 avec TPM virtuel.", "Algorithme documenté : XTS-AES 128.", "La clé de récupération est sauvegardée dans AD DS avant chiffrement.", "Contrôle : Get-BitLockerVolume, manage-bde -status et présence d’un objet de récupération AD — valeur masquée."])
code('''Get-BitLockerVolume -MountPoint "C:" | Select MountPoint,VolumeStatus,ProtectionStatus,EncryptionMethod
manage-bde -status C:
# À exécuter uniquement par un administrateur habilité ; ne pas exporter la clé dans le dossier.''')
proof("it-2/bitlocker OK.png", "État BitLocker sur POSTE-01", "Vérifier le chiffrement et l’activation de la protection du volume système.", "La capture atteste un état BitLocker opérationnel ; aucune clé n’est retranscrite ici.")
doc.add_heading("5.2 Windows LAPS", level=2)
table(["Paramètre", "Valeur documentée"], [("Sauvegarde", "Windows Server Active Directory (BackupDirectory=2)"), ("Cible", "OU=Ordinateurs"), ("Délégation", "Auto-écriture des attributs LAPS aux ordinateurs"), ("Accès au secret", "Administrateurs explicitement habilités uniquement")])
code('''Get-LapsADPassword -Identity "POSTE-01" -AsPlainText:$false
Get-WinEvent -LogName "Microsoft-Windows-LAPS/Operational" -MaxEvents 20 |
  Select TimeCreated, Id, LevelDisplayName, Message
# Ne jamais capturer ou exporter la valeur du mot de passe.''')
proof("it-2/LAPS OK Coté client.png", "Traitement LAPS côté client", "Contrôler que POSTE-01 traite la stratégie et renouvelle son secret local.", "Le traitement LAPS est visible côté client ; le secret n’est pas inclus.")
proof("it-4/Laps ok.png", "Récupération LAPS côté serveur — valeur protégée", "Vérifier que l’objet ordinateur publie les données LAPS dans AD et que l’accès est délégué, sans demander l’affichage en clair.", "La récupération a été testée depuis SRV-AD01 ; le champ Password reste un objet SecureString et aucun secret n’est exposé.")

doc.add_heading("6. SRV-FIC01, partages et modèle AGDLP", level=1)
table(["Composant", "Configuration"], [("VM", "2 vCPU, 4 Go RAM, système 60 Go"), ("Données", "D: DATA, 40 Go ; D:\\DATA"), ("Sauvegarde", "E: BACKUP, 40 Go en laboratoire"), ("Rôle", "FS-FileServer"), ("Domaine / OU", "corp.local / OU=Serveurs")])
table(["Ressource", "UNC", "Groupe global → local", "NTFS / SMB", "Lecteur"], [
    ("RH", "\\\\SRV-FIC01\\RH", "GG_RH → DL_RH_RW", "Modify / Change", "H:"),
    ("IT", "\\\\SRV-FIC01\\IT", "GG_IT → DL_IT_RW", "Modify / Change", "I:"),
    ("COMMUN", "\\\\SRV-FIC01\\COMMUN", "GG_RH + GG_IT → DL_COMMUN_RW", "Modify / Change", "S:"),
])
doc.add_paragraph("AGDLP : Account → Global → Domain Local → Permission. Les identités sont affectées aux groupes globaux métier ; ceux-ci sont imbriqués dans les groupes locaux de domaine qui reçoivent les ACL. Cette séparation facilite les arrivées, départs et audits.")
code('''Get-SmbShare | Select Name,Path,Description
Get-SmbShareAccess -Name RH,IT,COMMUN
icacls D:\\DATA\\RH
icacls D:\\DATA\\IT
icacls D:\\DATA\\COMMUN
Get-ADGroupMember GG_RH
Get-ADGroupMember DL_RH_RW''')
proof("it-3/lecteurs-reseau-it.png", "Lecteurs réseau d’un compte IT", "Valider le mappage conditionnel des lecteurs par appartenance aux groupes.", "Le poste IT obtient les lecteurs correspondant à ses habilitations.")
proof("it-3/vue-acces-geres.png", "Accès aux ressources gérés", "Contrôler que les partages et leurs permissions sont accessibles selon le modèle AGDLP.", "Les accès métier sont gérés par groupes et non par ACL utilisateur directe.")

doc.add_heading("7. Shadow Copies / Versions précédentes", level=1)
doc.add_paragraph("Les clichés VSS sont configurés sur D: afin de restaurer rapidement une version antérieure. Ils ne remplacent pas une sauvegarde : la perte du volume D: peut détruire données et clichés.")
code('''vssadmin list shadowstorage /for=D:
vssadmin list shadows /for=D:
# Configuration laboratoire si absente :
vssadmin add shadowstorage /for=D: /on=D: /maxsize=10GB
vssadmin create shadow /for=D:''')
table(["Exploitation", "Consigne"], [("Planification", "Créer plusieurs points selon la fréquence de modification et la capacité disponible"), ("Surveillance", "Contrôler l’espace shadowstorage et l’ancienneté du plus vieux cliché"), ("Restauration", "Depuis Propriétés > Versions précédentes ; préférer Copier avant Restaurer"), ("Traçabilité", "Noter fichier, demandeur, version, destination et résultat")])
proof("it-3/cliche-instant-ok.png", "Cliché VSS créé", "Vérifier la présence d’un point de restauration sur D:.", "Un cliché instantané est présent et exploitable pour les versions précédentes.")
proof("it-3/version-precedente-ok.png", "Restauration d’une version précédente", "Valider la récupération d’une version antérieure depuis un client autorisé.", "Le scénario de version précédente a abouti.")

doc.add_heading("8. Sauvegarde et restauration", level=1)
table(["Paramètre", "Valeur"], [("Serveur", "SRV-FIC01"), ("Outil", "Windows Server Backup"), ("Périmètre", "D:\\DATA"), ("Destination labo", "E:\\BACKUP"), ("Tests", "Version antérieure, fichier supprimé, dossier supprimé"), ("Limite", "E: sur le même hôte ne protège pas contre la perte totale ; externaliser en production")])
code('''Get-WindowsFeature Windows-Server-Backup
wbadmin start backup -backupTarget:E: -include:D:\\DATA -quiet
wbadmin get versions -backupTarget:E:
wbadmin get items -version:<VERSION> -backupTarget:E:
# Restaurer d'abord vers un emplacement alternatif, valider, puis remettre en production.''')
doc.add_heading("8.1 Procédure de restauration contrôlée", level=2)
for i, item in enumerate(["Qualifier la demande : objet, chemin, date/heure, propriétaire et urgence.", "Identifier la source la plus adaptée : VSS pour une version récente ; WSB pour suppression ou sinistre.", "Consigner la version retenue et vérifier la capacité de destination.", "Restaurer dans un dossier alternatif lorsque possible.", "Faire valider l’intégrité et les ACL par le propriétaire métier.", "Replacer en production, puis consigner durée, résultat et éventuels écarts."], 1): doc.add_paragraph(f"{i}. {item}")
proof("it-3/backup-once-ok.png", "Sauvegarde ponctuelle réussie", "Vérifier qu’une version de sauvegarde de D:\\DATA existe sur la cible E:.", "La tâche de sauvegarde ponctuelle a terminé avec succès.")
proof("it-3/it-dossier-restaure-ok.png", "Dossier IT restauré", "Tester un scénario de suppression puis restauration depuis Windows Server Backup.", "Le dossier supprimé a été récupéré, démontrant la restaurabilité de la sauvegarde.")

doc.add_heading("9. Scripts PowerShell d’exploitation", level=1)
doc.add_paragraph("Les scripts sources sont conservés dans docs/assets/files/admin-windows/it-4. Leur usage doit être précédé d’une revue, d’un test limité et, pour Create-Users.ps1, d’une simulation -WhatIf.")
table(["Script", "But", "Entrées / sorties", "Précaution"], [
    ("Create-Users.ps1", "Créer les comptes depuis CSV et les ajouter aux groupes", "Users-Source.csv → rapport CSV", "Mot de passe temporaire généré ; ne jamais journaliser sa valeur"),
    ("Inventory-AD.ps1", "Exporter utilisateurs, groupes, membres, ordinateurs et OU", "CSV + rapport HTML horodaté", "Protéger l’export ; il contient des données d’annuaire"),
    ("Commandes firewall", "Créer/exporter les règles par GPO", "PolicyStore → CSV", "Tester depuis POSTE-01 avant généralisation"),
])
code('''# Exécution conseillée
.\\Create-Users.ps1 -CsvPath .\\Users-Source.csv -WhatIf
.\\Inventory-AD.ps1 -GenerateHtml
Start-Transcript -Path C:\\Logs\\Exploitation-$(Get-Date -Format yyyyMMdd-HHmmss).txt
# ... opérations ...
Stop-Transcript''')
for source in [ROOT / "docs/assets/files/admin-windows/it-4/Create-Users.ps1", ROOT / "docs/assets/files/admin-windows/it-4/Inventory-AD.ps1"]:
    doc.add_heading(f"9.{1 if 'Create' in source.name else 2} Code intégral — {source.name}", level=2)
    code(source.read_text(encoding="utf-8-sig"))

doc.add_heading("10. Inventaire pare-feu", level=1)
doc.add_paragraph("Les deux profils Domaine sont attendus actifs avec trafic entrant bloqué par défaut et sortant autorisé. Les CSV présents prouvent les règles explicitement autorisées ci-dessous.")
def load_csv(name):
    with (ROOT / "docs/assets/files/admin-windows/it-4" / name).open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))
for title, filename in [("SRV-AD01", "GPO-FW-SRV-AD01-Rules.csv"), ("SRV-FIC01", "GPO-FW-SRV-FIC01-Rules.csv")]:
    doc.add_heading(f"10.{1 if title=='SRV-AD01' else 2} {title}", level=2)
    rows = load_csv(filename)
    table(["Règle", "État", "Direction", "Action", "Profil"], [(x["DisplayName"], x["Enabled"], x["Direction"], x["Action"], x["Profile"]) for x in rows])
doc.add_paragraph("Portée RDP : la procédure limite 3389 au réseau d’administration 10.42.0.0/24. Point à revalider : l’export CSV ne contient pas RemoteAddress ; contrôler la règle effective avant reprise.")
table(["Service", "Serveur", "Ports", "Justification"], [
    ("DNS", "SRV-AD01", "53 TCP/UDP", "Résolution AD"), ("Kerberos", "SRV-AD01", "88 TCP/UDP", "Authentification"),
    ("LDAP/LDAPS", "SRV-AD01", "389 TCP/UDP, 636 TCP", "Annuaire ; 636 seulement si certificat et usage validés"),
    ("Catalogue global", "SRV-AD01", "3268-3269 TCP", "Recherche forêt"), ("SMB", "SRV-AD01/SRV-FIC01", "445 TCP", "SYSVOL/NETLOGON et fichiers"),
    ("RPC", "SRV-AD01/SRV-FIC01", "135 TCP + ports dynamiques à valider", "Administration et services Windows"),
    ("RDP", "SRV-AD01/SRV-FIC01", "3389 TCP", "Administration restreinte au réseau admin"),
])
code('''Get-NetFirewallProfile | Select Name,Enabled,DefaultInboundAction,DefaultOutboundAction
Get-NetFirewallRule -PolicyStore "corp.local\\GPO-FW-SRV-AD01" | Get-NetFirewallPortFilter
Test-NetConnection SRV-AD01 -Port 53
Test-NetConnection SRV-FIC01 -Port 445
# Tester aussi les ports qui doivent rester fermés (80, 21, etc.).''')
proof("it-4/test SRV-FIC01.png", "Tests de flux vers SRV-FIC01", "Valider depuis POSTE-01 que SMB et les seuls flux d’administration justifiés sont accessibles.", "Les tests réseau du serveur de fichiers ont été consignés.")

doc.add_heading("11. Supervision, maintenance et reprise", level=1)
table(["Fréquence", "Contrôles"], [
    ("Quotidien", "Services AD/DNS/SMB, espace D:/E:, dernières sauvegardes, événements critiques"),
    ("Hebdomadaire", "Échecs 4625, état VSS, versions WSB, réplication/SYSVOL, changements de groupes privilégiés"),
    ("Mensuel", "Test de restauration échantillonné, revue ACL/AGDLP, inventaire GPO/firewall, capacité des volumes"),
    ("Trimestriel", "Test complet de reprise, revue des comptes et délégations, mise à jour du dossier"),
])
code('''dcdiag /v
Get-Service DNS,NTDS,LanmanServer
Get-Volume | Select DriveLetter,FileSystemLabel,SizeRemaining,Size
wbadmin get versions -backupTarget:E:
Get-WinEvent -FilterHashtable @{LogName='Security';Id=4625;StartTime=(Get-Date).AddDays(-1)}
Get-WinEvent -FilterHashtable @{LogName='System';Level=1,2;StartTime=(Get-Date).AddDays(-1)}''')
doc.add_heading("11.1 Ordre de reprise recommandé", level=2)
bullets(["Rétablir le réseau, la passerelle et l’hôte LABO/Hyper-V.", "Démarrer et valider SRV-AD01 : IP, DNS, AD DS, SYSVOL, NETLOGON et authentification.", "Démarrer SRV-FIC01 : jointure domaine, volumes D:/E:, rôle fichiers, partages et ACL.", "Démarrer POSTE-01, forcer les GPO et valider connexion, lecteurs, BitLocker et LAPS.", "Tester les sauvegardes et les flux pare-feu ; documenter tout écart."])

doc.add_heading("12. Points ouverts et critères de passation", level=1)
table(["Point", "Risque", "Action de reprise"], [
    ("LDAPS 636 ouvert", "Service inutile si aucun certificat/usage", "Valider le certificat et l’usage, sinon retirer l’autorisation"),
    ("RPC dynamique non décrit dans l’export", "Administration distante partielle", "Tester les outils réellement utilisés et documenter la plage requise"),
    ("Sauvegarde E: sur le même hôte", "Perte commune serveur/disques", "Ajouter une copie externalisée/offline selon 3-2-1"),
    ("Pas d’export exhaustif des paramètres GPO", "Difficile de comparer la dérive", "Générer Get-GPOReport -ReportType Html pour chaque GPO"),
])
bullets(["L’administrateur repreneur sait accéder aux consoles sans recevoir de secret dans ce document.", "Il sait retrouver les sauvegardes, effectuer une restauration alternative et valider les ACL.", "Il sait expliquer le modèle AGDLP et identifier les GPO appliquées.", "Il sait contrôler BitLocker/LAPS sans afficher de secret.", "Il sait reproduire l’inventaire AD et pare-feu et interpréter les écarts."])

doc.add_heading("Annexe A — Emplacement des preuves sources", level=1)
table(["Type", "Emplacement dans le dépôt"], [
    ("Captures Windows", "docs/assets/img/admin-windows/it-1 à it-4"),
    ("Scripts et exports", "docs/assets/files/admin-windows/it-4"),
    ("Inventaire AD HTML", "docs/admin-windows/it-4/AD-Inventory.html"),
    ("Procédures", "docs/admin-windows/it-1 à it-4"),
])
doc.add_paragraph("Fin du dossier — toute modification d’infrastructure doit entraîner une mise à jour de la version, de la date et de la section concernée.")

doc.save(OUT)
CONVENTION.write_bytes(OUT.read_bytes())
print(OUT)
print(CONVENTION)
